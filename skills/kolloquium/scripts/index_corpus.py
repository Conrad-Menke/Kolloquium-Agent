#!/usr/bin/env python3
"""PDFs und DOCX-Dateien in einen lokalen Chroma-Vektor-Store indizieren.

Akzeptiert entweder eine einzelne Datei oder ein Verzeichnis (rekursiv nach
unterstützten Formaten durchsucht). Jeder Chunk speichert seinen
Quelldatei-Pfad und seine Seitenzahl (falls verfügbar), damit der
Prüfer-Agent die exakte Passage zitieren kann, auf der eine Frage fundiert.

Unterstützte Formate:
    .pdf   — Seitenzahl bleibt erhalten (1-indiziert)
    .docx  — keine nativen Seiten; page ist None, Zitation nur nach Dateiname

Aufruf:
    python index_corpus.py <pfad> [--index-dir VERZ] [--chunk-size N] [--overlap N]

<pfad>  Eine unterstützte Datei oder ein Ordner, der rekursiv nach
        unterstützten Dateien durchsucht wird.

Die Ausgabe wird in die Chroma-Collection `kolloquium_passages` innerhalb des
Index-Verzeichnisses geschrieben (Default: ../index relativ zu diesem Skript).

Exit-Codes:
    0  mindestens eine Datei indiziert (oder bereits aktuell)
    1  Pfad nicht gefunden / keine unterstützten Dateien gefunden
    2  eine oder mehrere Dateien hatten keinen extrahierbaren Text
"""
from __future__ import annotations

import argparse
import hashlib
import sys
from pathlib import Path

import chromadb
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_INDEX_DIR = SCRIPT_DIR.parent / "index"
from common import COLLECTION_NAME, EMBED_MODEL
SUPPORTED_EXTS = {".pdf", ".docx"}


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("path", type=Path, help="Datei oder Ordner zum Indizieren")
    p.add_argument(
        "--index-dir",
        type=Path,
        default=DEFAULT_INDEX_DIR,
        help=f"Chroma-DB-Verzeichnis (Default: {DEFAULT_INDEX_DIR})",
    )
    p.add_argument("--chunk-size", type=int, default=500, help="Zeichen pro Chunk")
    p.add_argument("--overlap", type=int, default=80, help="Überlappung zwischen Chunks (Zeichen)")
    p.add_argument("--force", action="store_true", help="Neu indizieren, auch wenn Datei schon indiziert")
    return p.parse_args()


def find_sources(path: Path) -> list[Path]:
    """Unterstützte Dateien zurückgeben. Akzeptiert eine einzelne Datei oder
    ein Verzeichnis (rekursiv)."""
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_EXTS else []
    if path.is_dir():
        return sorted(p for p in path.rglob("*") if p.suffix.lower() in SUPPORTED_EXTS)
    return []


def extract_pages_pdf(pdf_path: Path) -> list[tuple[int, str]]:
    """Liste von (seitenzahl, text) zurückgeben. Seitenzahlen 1-indiziert."""
    reader = PdfReader(str(pdf_path))
    pages: list[tuple[int, str]] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((i, text))
    return pages


def extract_text_docx(docx_path: Path) -> list[tuple[int | None, str]]:
    """Ganze DOCX als einzelnen logischen Block zurückgeben. DOCX hat keine
    nativen Seiten, daher page=None — Zitation nur nach Dateiname.
    Tabellen-Zellen werden ebenfalls erfasst (pro Zeile dedupliziert, damit
    verbundene Zellen nicht mehrfach indiziert werden)."""
    from docx import Document

    doc = Document(str(docx_path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            seen: set[str] = set()
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    parts.append(t)
    text = "\n".join(parts)
    return [(None, text)] if text else []


def extract(file_path: Path) -> list[tuple[int | None, str]]:
    """Dispatch nach Erweiterung. Gibt Liste von (seite|None, text) zurück."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_pages_pdf(file_path)
    if ext == ".docx":
        return extract_text_docx(file_path)
    return []


def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    """Gieriger Chunker fester Größe mit Überlappung. Größe in Zeichen."""
    if size <= 0:
        raise ValueError(f"chunk-size ({size}) muss positiv sein")
    if overlap < 0:
        raise ValueError(f"overlap ({overlap}) darf nicht negativ sein")
    if overlap >= size:
        raise ValueError(f"overlap ({overlap}) muss kleiner als chunk-size ({size}) sein")
    if len(text) <= size:
        return [text] if text else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks


def file_signature(file_path: Path) -> str:
    """Stabile Signatur für Dedup: Pfad + mtime + Größe."""
    stat = file_path.stat()
    raw = f"{file_path.resolve()}|{stat.st_mtime_ns}|{stat.st_size}".encode()
    return hashlib.sha1(raw).hexdigest()  # nosec: Dedup-Key, keine Sicherheit


def index_one(
    file_path: Path,
    embedder: SentenceTransformer,
    collection,
    chunk_size: int,
    overlap: int,
    force: bool,
) -> tuple[bool, str]:
    """Eine einzelne Datei indizieren. Gibt (ok, meldung) zurück."""
    sig = file_signature(file_path)
    if not force:
        existing = collection.get(where={"source_sig": sig})
        if existing["ids"]:
            return True, f"bereits indiziert: {file_path.name} ({len(existing['ids'])} Chunks), übersprungen"

    try:
        blocks = extract(file_path)
    except Exception as exc:
        return False, f"Fehler beim Lesen von {file_path.name}: {exc}"
    if not blocks:
        return False, f"kein extrahierbarer Text in {file_path.name} (eingescanntes PDF braucht OCR)"

    docs: list[str] = []
    metas: list[dict] = []
    ids: list[str] = []
    for block_index, (page_no, text) in enumerate(blocks):
        for j, chunk in enumerate(chunk_text(text, chunk_size, overlap)):
            if not chunk.strip():
                continue
            docs.append(chunk)
            meta: dict = {
                "source_file": str(file_path.resolve()),
                "source_name": file_path.name,
                "source_ext": file_path.suffix.lower().lstrip("."),
                "block_index": block_index,
                "chunk_index": j,
                "source_sig": sig,
            }
            if page_no is not None:
                meta["page"] = page_no
            metas.append(meta)
            page_part = f"p{page_no}" if page_no is not None else f"b{block_index}"
            ids.append(f"{sig}::{page_part}::c{j}")

    if not docs:
        return False, f"alle Chunks leer in {file_path.name}"

    collection.delete(where={"source_file": str(file_path.resolve())})
    vectors = embedder.encode(docs, show_progress_bar=False, convert_to_numpy=True, normalize_embeddings=True).tolist()
    collection.upsert(ids=ids, documents=docs, metadatas=metas, embeddings=vectors)
    return True, f"{len(docs)} Chunks aus {file_path.name} indiziert"


def main() -> int:
    args = parse_args()
    path: Path = args.path
    if not path.exists():
        print(f"Fehler: Pfad nicht gefunden: {path}", file=sys.stderr)
        return 1

    sources = find_sources(path)
    if not sources:
        print(
            f"Fehler: keine unterstützten Dateien ({', '.join(sorted(SUPPORTED_EXTS))}) unter {path} gefunden",
            file=sys.stderr,
        )
        return 1

    print(f"{len(sources)} Datei(en) unter {path} gefunden")
    args.index_dir.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(path=str(args.index_dir))
    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"description": "Kolloquiums-Prüfer-Passagen"},
    )

    print(f"lade Embedder: {EMBED_MODEL}")
    embedder = SentenceTransformer(EMBED_MODEL)

    ok_count = 0
    fail_count = 0
    had_text_error = False
    for file_path in sources:
        ok, msg = index_one(
            file_path, embedder, collection, args.chunk_size, args.overlap, args.force
        )
        print(f"  - {msg}")
        if ok:
            ok_count += 1
        else:
            fail_count += 1
            if "OCR" in msg or "kein extrahierbarer Text" in msg:
                had_text_error = True

    total_chunks = collection.count()
    print(f"fertig: {ok_count} ok, {fail_count} fehlgeschlagen. Chunks im Index gesamt: {total_chunks}")
    if had_text_error:
        return 2
    return 0 if ok_count > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
